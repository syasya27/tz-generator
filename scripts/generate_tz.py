import os
import asyncio
import pandas as pd
from openai import AsyncOpenAI
from docx import Document

# API
client = AsyncOpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Данные
df = pd.read_csv("data/tz_input.csv")

with open("prompts/tz_prompt.txt", encoding="utf-8") as f:
    prompt_template = f.read()

# Папка
os.makedirs("output", exist_ok=True)

# Ограничение параллельности (ВАЖНО)
SEM = asyncio.Semaphore(5)


# 📄 Сохранение в DOCX
def save_as_docx(content, filename):
    doc = Document()

    for line in content.split("\n"):
        line = line.strip()

        if line.startswith("# "):
            doc.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("### "):
            doc.add_heading(line.replace("### ", ""), level=3)
        elif line == "":
            doc.add_paragraph("")  # пустая строка
        else:
            doc.add_paragraph(line)

    doc.save(filename)


async def generate_tz(i, row):
    async with SEM:  # контроль параллельности

        competitors = "\n".join([
            str(row.get("r1", "")),
            str(row.get("r2", "")),
            str(row.get("r3", ""))
        ])

        prompt = prompt_template.format(
            url=row.get("url", ""),
            query=row.get("query", ""),
            format=row.get("format", ""),
            competitors=competitors,
            volume=row.get("Current volume", "")
        )

        try:
            response = await client.responses.create(
                model="gpt-4.1-mini",
                input=prompt
            )

            text = response.output_text

            filename = f"output/tz_{i+1}.docx"
            save_as_docx(text, filename)

            print(f"✅ Done: {filename}")

        except Exception as e:
            print(f"❌ Error at row {i}: {e}")


async def main():
    tasks = []

    for i, row in df.iterrows():
        tasks.append(generate_tz(i, row))

    await asyncio.gather(*tasks)


if __name__ == "__main__":
    asyncio.run(main())
